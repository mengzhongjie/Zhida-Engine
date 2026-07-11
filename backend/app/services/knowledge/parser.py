"""
智答引擎（ZhiDa Engine）—— 文档解析器工厂

支持多格式文档解析，含降级策略：
- PDF: pdfplumber（表格）+ unstructured（文本）
- Word: python-docx
- Excel: openpyxl + pandas
- TXT/Markdown: 原生读取

降级策略：
- 内存不足 → 纯文本提取（放弃表格和结构）
- 文件过大 → 分页解析
- 解析超时 → 跳过当前页，继续下一页
"""

import os
import time
import asyncio
from pathlib import Path
from typing import Optional, AsyncIterator
from dataclasses import dataclass, field
from enum import Enum

from loguru import logger


class FileType(str, Enum):
    """支持的文件类型"""
    PDF = "pdf"
    DOCX = "docx"
    XLSX = "xlsx"
    TXT = "txt"
    MD = "md"
    CSV = "csv"
    JSON = "json"
    XML = "xml"


class ParseStatus(str, Enum):
    """解析状态"""
    SUCCESS = "success"
    DEGRADED = "degraded"  # 降级解析
    PARTIAL = "partial"    # 部分解析
    FAILED = "failed"


@dataclass
class ParseResult:
    """解析结果"""
    status: ParseStatus
    text: str = ""
    pages: list[str] = field(default_factory=list)  # 分页文本
    metadata: dict = field(default_factory=dict)     # 元数据（标题、页码等）
    tables: list[str] = field(default_factory=list)  # 表格数据（Markdown 格式）
    error_message: str = ""
    parse_time_ms: float = 0.0


@dataclass
class ParseChunk:
    """解析文本块 —— 用于后续切片"""
    text: str
    metadata: dict = field(default_factory=dict)
    chunk_index: int = 0


class DocumentParser:
    """
    文档解析器 —— 工厂模式，根据文件类型选择解析策略

    Usage:
        parser = DocumentParser()

        # 完整解析
        result = await parser.parse("document.pdf")

        # 分页解析（大文件）
        async for chunk in parser.parse_paginated("large_doc.pdf", page_size=5):
            process(chunk)
    """

    # 文件大小阈值
    MAX_FILE_SIZE_MB = 100  # 超过 100MB 触发分页解析
    PARSE_TIMEOUT_SEC = 300  # 解析超时（秒）

    # MinerU 专有格式（现有 FileType 枚举未覆盖的格式）
    # 这些格式必须启用 MinerU 才能解析
    MINERU_ONLY_FORMATS = {
        "pptx", "ppt", "epub", "html", "htm",
        "png", "jpg", "jpeg", "bmp", "tiff", "webp",
    }

    def __init__(self, mineru_parser: Optional["MinerUParser"] = None):
        self._supported_formats = {
            FileType.PDF: self._parse_pdf,
            FileType.DOCX: self._parse_docx,
            FileType.XLSX: self._parse_xlsx,
            FileType.TXT: self._parse_txt,
            FileType.MD: self._parse_txt,  # Markdown 使用文本解析
            FileType.CSV: self._parse_csv,
            FileType.JSON: self._parse_txt,  # JSON 使用文本解析
            FileType.XML: self._parse_txt,   # XML 使用文本解析
        }

        # MinerU 解析器（可选注入，支持依赖注入便于测试）
        if mineru_parser is not None:
            self._mineru = mineru_parser
        else:
            self._mineru = self._create_mineru_if_enabled()

    @staticmethod
    def _create_mineru_if_enabled():
        """延迟创建 MinerU 实例，避免导入错误阻塞应用启动

        仅当 ENABLE_MINERU=True 时尝试加载 MinerU 模块。
        """
        try:
            from app.services.knowledge.mineru.parser import create_mineru_parser
            return create_mineru_parser()
        except ImportError:
            return None
        except Exception as e:
            logger.warning(f"无法加载 MinerU，将使用本地解析器: {e}")
            return None

    def get_file_type(self, file_path: str) -> Optional[FileType]:
        """根据文件扩展名获取文件类型

        Returns:
            FileType 枚举，如果扩展名不在支持列表中则返回 None
            （None 可能是 MinerU 专有格式，需由 MinerU 处理）
        """
        ext = Path(file_path).suffix.lower().lstrip(".")
        try:
            return FileType(ext)
        except ValueError:
            return None

    def is_mineru_only_format(self, file_path: str) -> bool:
        """判断是否仅 MinerU 支持的格式"""
        ext = Path(file_path).suffix.lower().lstrip(".")
        return ext in self.MINERU_ONLY_FORMATS

    async def parse(self, file_path: str) -> ParseResult:
        """
        解析文档 —— 自动选择解析策略，含降级

        解析优先级：
        1. MinerU 解析（若启用且文件格式匹配）
        2. 现有解析器（pdfplumber/python-docx/openpyxl 等）
        3. 纯文本降级（内存不足时）
        4. 分页解析（超时或文件过大时）

        Args:
            file_path: 文件路径

        Returns:
            ParseResult 包含解析后的文本、元数据等
        """
        start_time = time.time()
        file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
        file_ext = Path(file_path).suffix.lower().lstrip(".")
        file_type = self.get_file_type(file_path)

        # ================================================================
        # 策略 1: MinerU 优先（若启用）
        # ================================================================
        if self._mineru and self._mineru.can_handle(file_ext):
            file_name = Path(file_path).name

            # 检查文件大小是否超过 MinerU 阈值
            if file_size_mb > self._mineru.config.max_file_size_mb:
                logger.info(f"MinerU 跳过（文件过大 {file_size_mb:.0f}MB > {self._mineru.config.max_file_size_mb}MB）, 使用本地解析器: {file_name}")
            elif not await self._mineru.is_available():
                logger.info(f"MinerU 不可用，使用本地解析器: {file_name}")
            else:
                try:
                    result = await asyncio.wait_for(
                        self._mineru.parse(file_path),
                        timeout=self.PARSE_TIMEOUT_SEC,
                    )
                    result.parse_time_ms = (time.time() - start_time) * 1000

                    if result.status != ParseStatus.FAILED and result.text.strip():
                        logger.info(f"MinerU 解析成功: {file_name}, 耗时={result.parse_time_ms:.0f}ms")
                        return result

                    logger.warning(f"MinerU 返回空或失败结果，降级到本地解析器: {file_name}")
                except asyncio.TimeoutError:
                    logger.warning(f"MinerU 解析超时 ({self.PARSE_TIMEOUT_SEC}s)，降级到本地解析器: {file_name}")
                except Exception as e:
                    logger.warning(f"MinerU 解析失败，降级到本地解析器: {file_name}: {e}")

                if self._mineru.config.fallback_on_failure:
                    # 继续到策略 2
                    pass
                else:
                    # 不降级，直接报错
                    logger.error(f"MinerU 解析失败且 fallback_on_failure=False: {file_name}")
                    return ParseResult(
                        status=ParseStatus.FAILED,
                        error_message=f"MinerU 解析失败: 等待重试",
                    )

        # ================================================================
        # 策略 2: 现有解析器（或 MinerU 专有格式检查）
        # ================================================================
        # MinerU 专有格式必须启用 MinerU
        if file_type is None:
            if self.is_mineru_only_format(file_path):
                error_msg = (
                    f"格式 .{file_ext} 需要启用 MinerU 才能解析。"
                    f"请设置环境变量 ZHIDA_ENABLE_MINERU=true 并安装 magic-pdf"
                )
                logger.error(error_msg)
                return ParseResult(
                    status=ParseStatus.FAILED,
                    error_message=error_msg,
                )
            return ParseResult(
                status=ParseStatus.FAILED,
                error_message=f"不支持的文件类型: .{file_ext}",
            )

        logger.info(f"开始本地解析: {Path(file_path).name} (类型={file_type.value}, 大小={file_size_mb:.1f}MB)")

        try:
            # 大文件使用分页解析
            if file_size_mb > self.MAX_FILE_SIZE_MB:
                logger.warning(f"文件过大 ({file_size_mb:.1f}MB)，使用分页解析")
                return await self._parse_large_file(file_path, file_type)

            # 正常解析
            parser_fn = self._supported_formats.get(file_type)
            if parser_fn is None:
                raise ValueError(f"不支持的解析器: {file_type}")

            result = await self._parse_with_timeout(parser_fn, file_path)

        except MemoryError:
            logger.warning("内存不足，降级为纯文本提取")
            result = await self._degraded_text_only(file_path, file_type)

        except asyncio.TimeoutError:
            logger.warning("解析超时，降级为分页解析")
            result = await self._parse_large_file(file_path, file_type, page_size=5)

        except Exception as e:
            logger.error(f"解析失败: {file_path}: {e}")
            result = ParseResult(
                status=ParseStatus.FAILED,
                error_message=str(e),
            )

        result.parse_time_ms = (time.time() - start_time) * 1000
        logger.info(f"解析完成: {Path(file_path).name}, 状态={result.status.value}, 耗时={result.parse_time_ms:.0f}ms, 文本长度={len(result.text)}")

        return result

    async def parse_paginated(
        self,
        file_path: str,
        page_size: int = 5,
    ) -> AsyncIterator[ParseChunk]:
        """
        分页解析文档 —— 用于大文件，每批处理 N 页
        优先使用 MinerU 分页（若启用），否则使用现有解析器。

        Usage:
            async for chunk in parser.parse_paginated("large.pdf"):
                await process_chunk(chunk)
        """
        file_ext = Path(file_path).suffix.lower().lstrip(".")

        # MinerU 分页（若启用且支持该格式）
        if self._mineru and self._mineru.can_handle(file_ext):
            try:
                async for chunk in self._mineru.parse_paginated(file_path, page_size):
                    yield chunk
                return
            except Exception as e:
                logger.warning(f"MinerU 分页解析失败，降级: {e}")
                if self._mineru and not self._mineru.config.fallback_on_failure:
                    raise

        # 现有解析器分页
        result = await self.parse(file_path)

        if result.status == ParseStatus.FAILED:
            logger.error(f"分页解析失败: {result.error_message}")
            return

        # 按页输出
        for i, page_text in enumerate(result.pages):
            yield ParseChunk(
                text=page_text,
                metadata={
                    **result.metadata,
                    "page": i + 1,
                    "total_pages": len(result.pages),
                },
                chunk_index=i,
            )

    async def _parse_with_timeout(self, fn, *args) -> ParseResult:
        """带超时的解析"""
        try:
            return await asyncio.wait_for(
                asyncio.to_thread(fn, *args),
                timeout=self.PARSE_TIMEOUT_SEC,
            )
        except asyncio.TimeoutError:
            raise

    # ================================================================
    # 各格式解析器
    # ================================================================

    def _parse_pdf(self, file_path: str) -> ParseResult:
        """解析 PDF 文档 —— 使用 pdfplumber 提取文本和表格"""
        import pdfplumber

        pages = []
        tables = []
        metadata = {}

        with pdfplumber.open(file_path) as pdf:
            metadata = dict(pdf.metadata or {})

            for i, page in enumerate(pdf.pages):
                # 提取文本
                text = page.extract_text() or ""
                pages.append(text)

                # 提取表格
                page_tables = page.extract_tables()
                for table in page_tables:
                    if table:
                        table_md = self._table_to_markdown(table)
                        tables.append(table_md)

        full_text = "\n\n".join(pages)

        return ParseResult(
            status=ParseStatus.SUCCESS,
            text=full_text,
            pages=pages,
            metadata={
                **metadata,
                "total_pages": len(pages),
                "table_count": len(tables),
            },
            tables=tables,
        )

    def _parse_docx(self, file_path: str) -> ParseResult:
        """解析 Word 文档 —— 使用 python-docx"""
        from docx import Document

        doc = Document(file_path)
        paragraphs = []

        for para in doc.paragraphs:
            if para.text.strip():
                # 保留段落样式信息
                if para.style.name.startswith("Heading"):
                    paragraphs.append(f"## {para.text}")
                else:
                    paragraphs.append(para.text)

        # 提取表格
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                cells = [cell.text for cell in row.cells]
                rows.append(" | ".join(cells))
            if rows:
                tables.append("\n".join(rows))

        full_text = "\n\n".join(paragraphs)

        return ParseResult(
            status=ParseStatus.SUCCESS,
            text=full_text,
            pages=[full_text],
            metadata={
                "paragraph_count": len(paragraphs),
                "table_count": len(tables),
            },
            tables=tables,
        )

    def _parse_xlsx(self, file_path: str) -> ParseResult:
        """解析 Excel 文档 —— 使用 openpyxl + pandas"""
        import pandas as pd

        # 读取所有 Sheet
        xlsx = pd.ExcelFile(file_path)
        sheets = []
        tables = []

        for sheet_name in xlsx.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            if df.empty:
                continue

            # 转 Markdown 表格
            table_md = df.to_markdown(index=False)
            sheets.append(f"### Sheet: {sheet_name}\n{table_md}")
            tables.append(table_md)

        full_text = "\n\n".join(sheets)

        return ParseResult(
            status=ParseStatus.SUCCESS,
            text=full_text,
            pages=sheets,
            metadata={
                "sheet_count": len(sheets),
                "total_rows": sum(len(t.split("\n")) for t in sheets),
            },
            tables=tables,
        )

    def _parse_txt(self, file_path: str) -> ParseResult:
        """解析纯文本/Markdown 文档"""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()

        return ParseResult(
            status=ParseStatus.SUCCESS,
            text=text,
            pages=[text],
            metadata={
                "char_count": len(text),
                "line_count": text.count("\n"),
            },
        )

    def _parse_csv(self, file_path: str) -> ParseResult:
        """解析 CSV 文档"""
        import pandas as pd

        df = pd.read_csv(file_path)

        if df.empty:
            return ParseResult(
                status=ParseStatus.SUCCESS,
                text="",
                metadata={"row_count": 0},
            )

        table_md = df.to_markdown(index=False)

        return ParseResult(
            status=ParseStatus.SUCCESS,
            text=table_md,
            tables=[table_md],
            metadata={
                "row_count": len(df),
                "column_count": len(df.columns),
            },
        )

    # ================================================================
    # 降级策略
    # ================================================================

    async def _parse_large_file(self, file_path: str, file_type: FileType, page_size: int = 5) -> ParseResult:
        """分页解析大文件 —— 降级策略"""
        if file_type == FileType.PDF:
            return await self._parse_pdf_paginated(file_path, page_size)
        else:
            # 非 PDF 文件，按行分页
            return self._parse_text_paginated(file_path, page_size)

    async def _parse_pdf_paginated(self, file_path: str, page_size: int = 5) -> ParseResult:
        """分页解析 PDF"""
        import pdfplumber

        pages = []
        tables = []

        with pdfplumber.open(file_path) as pdf:
            total_pages = len(pdf.pages)

            for i in range(0, total_pages, page_size):
                batch_pages = pdf.pages[i : i + page_size]
                batch_text = []

                for page in batch_pages:
                    try:
                        text = page.extract_text() or ""
                        batch_text.append(text)
                    except Exception as e:
                        logger.warning(f"页面 {page.page_number} 解析失败: {e}")
                        batch_text.append(f"[页面 {page.page_number} 解析失败]")

                pages.append("\n\n".join(batch_text))

        full_text = "\n\n".join(pages)

        return ParseResult(
            status=ParseStatus.PARTIAL,
            text=full_text,
            pages=pages,
            metadata={"total_pages": len(pages), "parse_mode": "paginated"},
        )

    def _parse_text_paginated(self, file_path: str, page_size: int = 5) -> ParseResult:
        """按行分页解析文本文件"""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            lines = f.readlines()

        pages = []
        for i in range(0, len(lines), page_size * 50):  # 每页约 50 行
            batch = lines[i : i + page_size * 50]
            pages.append("".join(batch))

        full_text = "\n\n".join(pages)

        return ParseResult(
            status=ParseStatus.PARTIAL,
            text=full_text,
            pages=pages,
            metadata={"total_pages": len(pages), "parse_mode": "paginated"},
        )

    async def _degraded_text_only(self, file_path: str, file_type: FileType) -> ParseResult:
        """
        纯文本降级解析 —— 放弃表格、结构，只提取文字

        触发条件：内存不足
        """
        logger.warning(f"降级为纯文本提取: {file_path}")

        if file_type == FileType.PDF:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                text = "\n".join([p.extract_text() or "" for p in pdf.pages])
        elif file_type == FileType.DOCX:
            from docx import Document
            doc = Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
        else:
            text = self._parse_txt(file_path).text

        return ParseResult(
            status=ParseStatus.DEGRADED,
            text=text,
            pages=[text],
            metadata={"parse_mode": "text_only_degraded"},
        )

    # ================================================================
    # 工具函数
    # ================================================================

    @staticmethod
    def _table_to_markdown(table: list[list]) -> str:
        """将二维数组转为 Markdown 表格"""
        if not table:
            return ""

        lines = []
        # 表头
        header = table[0]
        lines.append("| " + " | ".join(str(c or "") for c in header) + " |")
        lines.append("| " + " | ".join("---" for _ in header) + " |")

        # 数据行
        for row in table[1:]:
            lines.append("| " + " | ".join(str(c or "") for c in row) + " |")

        return "\n".join(lines)


# 全局解析器实例
document_parser = DocumentParser()